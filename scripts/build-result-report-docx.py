#!/usr/bin/env python3
"""Build Reqover's 2026 OSS contest report from the official DOCX template.

The retained template is never modified. The generated copy removes the guide
page and the non-applicable AI-model appendix, fills the official report slots,
and expands Attachment 1 from a CycloneDX JSON SBOM.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
from copy import deepcopy
from pathlib import Path
from urllib.parse import quote

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_NAME = "2026 오픈소스 개발자대회 결과보고서_접수번호(팀명).docx"
TEMPLATE_PATH = ROOT / "docs" / "competition" / "templates" / TEMPLATE_NAME
TEMPLATE_SHA256 = "937679bac40cbfaced3457530c232c9d190a74f6b5d67c58b4bc33014a579195"
DEFAULT_SBOM_PATH = ROOT / "sbom" / "reqover.cdx.json"
DEFAULT_OUTPUT_PATH = ROOT / "docs" / "competition" / "submission" / "Reqover_result_report_draft.docx"

FONT_NAME = "Malgun Gothic"
BODY_FONT_SIZE = Pt(10)
SMALL_FONT_SIZE = Pt(8)
SBOM_FONT_SIZE = Pt(7)
PLACEHOLDER_PREFIX = "[확인 필요:"

ASSET_MVC = ROOT / "docs" / "assets" / "reqover-mvc-request-attribution.png"
ASSET_REVERSE = ROOT / "docs" / "assets" / "reqover-code-to-endpoint-index.png"
ASSET_WEBFLUX = ROOT / "docs" / "assets" / "reqover-webflux-thread-hop.png"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--template", type=Path, default=TEMPLATE_PATH)
    parser.add_argument("--sbom", type=Path, default=DEFAULT_SBOM_PATH)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT_PATH)
    parser.add_argument("--team-name", default="[확인 필요: 참가 접수 정보와 동일한 팀명]")
    parser.add_argument("--team-size", default="[확인 필요: 팀장 포함 총 인원]")
    parser.add_argument("--division", default="[확인 필요: 학생 또는 일반]")
    parser.add_argument("--task-type", default="[확인 필요: 자유과제 또는 지정과제(기업명)]")
    parser.add_argument("--registration-number", default="[확인 필요: 접수번호]")
    parser.add_argument("--video-url", default="[확인 필요: 공개 또는 일부 공개 YouTube URL]")
    parser.add_argument(
        "--development-environment",
        default="[확인 필요: 최종 개발·검증 PC의 OS, CPU, RAM]",
    )
    parser.add_argument(
        "--team-contributions",
        default=(
            "[확인 필요: 팀원별 최종 역할·기여 확정] "
            "김태희—문제 정의와 제품 방향, core·instrumentation·Java Agent·report·sample 설계 및 통합. "
            "이상민—빌드·CI·크로스플랫폼 환경, core 안정화, Spring adapter·테스트·라이선스·공개 문서 정비."
        ),
    )
    parser.add_argument(
        "--verification",
        default=(
            "2026-08-14 JDK 17·21 clean build 35 tests 통과, CycloneDX 1.6 SBOM lock 일치, "
            "외부 Maven 구성요소 104개 OSV 0건 (log4j-api를 2.25.5로 올려 "
            "GHSA-qv9r-c865-cp47 해소). "
            "[확인 필요: v0.1.1 태그 커밋 해시와 Release CI 실행 URL — 태그 push 후 기입]. "
            "직전 기록: 2026-08-10 v0.1.0=62910fc1896d, "
            "https://github.com/reqover-labs/reqover/actions/runs/31366748139"
        ),
    )
    parser.add_argument(
        "--duplicate-benefit",
        default="[확인 필요: 정부지원·정부 대회 수상 등 중복수혜 여부와 별도 확인서 필요 여부]",
    )
    return parser.parse_args()


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def verify_template(path: Path) -> None:
    if not path.is_file():
        raise FileNotFoundError(f"Official template not found: {path}")
    actual = sha256(path)
    if actual != TEMPLATE_SHA256:
        raise ValueError(
            "Official template checksum mismatch. "
            f"expected={TEMPLATE_SHA256}, actual={actual}, path={path}"
        )


def set_run_font(
    run,
    *,
    size=BODY_FONT_SIZE,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str = "000000",
) -> None:
    run.font.name = FONT_NAME
    run.font.size = size
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attribute}"), FONT_NAME)
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "ko-KR")
    lang.set(qn("w:eastAsia"), "ko-KR")


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def format_paragraph(paragraph, *, after=2, before=0, line=1.05, alignment=None) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if alignment is not None:
        paragraph.alignment = alignment


def add_text_run(paragraph, text: str, *, bold=False, italic=False, size=BODY_FONT_SIZE, color="000000"):
    last_run = None
    for part in re.split(r"(\[확인 필요:[^\]]+\])", text):
        if not part:
            continue
        run = paragraph.add_run(part)
        set_run_font(run, size=size, bold=bold, italic=italic, color=color)
        if part.startswith(PLACEHOLDER_PREFIX) and part.endswith("]"):
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            run.font.color.rgb = RGBColor.from_string("C00000")
            run.bold = True
        last_run = run
    return last_run


def replace_paragraph(paragraph, text: str, *, bold=False, size=BODY_FONT_SIZE, alignment=None) -> None:
    clear_paragraph(paragraph)
    format_paragraph(paragraph, alignment=alignment)
    add_text_run(paragraph, text, bold=bold, size=size)


def clear_cell(cell) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def append_labeled_paragraph(cell, label: str, text: str, *, first=False, after=2) -> None:
    paragraph = cell.paragraphs[0] if first else cell.add_paragraph()
    if first:
        clear_paragraph(paragraph)
    format_paragraph(paragraph, after=after)
    if label:
        add_text_run(paragraph, f"{label}  ", bold=True, color="1F4E78")
    add_text_run(paragraph, text)


def set_cell_text(cell, text: str, *, bold=False, size=BODY_FONT_SIZE, center=False) -> None:
    clear_cell(cell)
    paragraph = cell.paragraphs[0]
    format_paragraph(
        paragraph,
        after=0,
        alignment=WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT,
    )
    add_text_run(paragraph, text, bold=bold, size=size)


def add_picture(cell, path: Path, caption: str, *, first=False, width=Inches(4.15)) -> None:
    if not path.is_file():
        paragraph = cell.paragraphs[0] if first else cell.add_paragraph()
        if first:
            clear_paragraph(paragraph)
        format_paragraph(paragraph, after=2)
        add_text_run(paragraph, f"[확인 필요: 이미지 파일 없음 — {path.name}]")
        return

    paragraph = cell.paragraphs[0] if first else cell.add_paragraph()
    if first:
        clear_paragraph(paragraph)
    format_paragraph(paragraph, after=1, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = paragraph.add_run()
    inline_shape = run.add_picture(str(path), width=width)
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", caption)
    doc_pr.set("descr", caption)

    caption_paragraph = cell.add_paragraph()
    format_paragraph(caption_paragraph, after=3, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text_run(caption_paragraph, caption, italic=True, size=SMALL_FONT_SIZE, color="595959")


def remove_element(element) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def remove_guide_and_ai_appendix(document: Document) -> None:
    if len(document.tables) < 9:
        raise ValueError("Unexpected official template structure: expected at least 9 tables")

    guide_table = document.tables[0]
    if "결과보고서 작성 안내" not in guide_table.cell(0, 0).text:
        raise ValueError("Unexpected official template structure: guide page table not found")
    remove_element(guide_table._element)

    sbom_note = next(
        (paragraph for paragraph in document.paragraphs if "필요 시, 행을 추가" in paragraph.text),
        None,
    )
    if sbom_note is None:
        raise ValueError("Unexpected official template structure: SBOM note not found")

    body = document.element.body
    note_index = list(body).index(sbom_note._p)
    for element in list(body)[note_index + 1 :]:
        if element.tag != qn("w:sectPr"):
            body.remove(element)


def paragraph_after(document: Document, element) -> object:
    body = document.element.body
    elements = list(body)
    index = elements.index(element)
    for candidate in elements[index + 1 :]:
        if candidate.tag == qn("w:p"):
            from docx.text.paragraph import Paragraph

            return Paragraph(candidate, document._body)
    raise ValueError("Expected paragraph not found after template element")


def fill_basic_information(document: Document, args: argparse.Namespace) -> None:
    # After guide removal, tables are: title, basic info, report body, SBOM title, SBOM table.
    if len(document.tables) != 5:
        raise ValueError(f"Unexpected generated structure: expected 5 tables, got {len(document.tables)}")
    info = document.tables[1]
    set_cell_text(info.rows[1].cells[1], args.team_name)
    set_cell_text(info.rows[1].cells[3], args.team_size)
    set_cell_text(info.rows[2].cells[1], args.division)
    set_cell_text(info.rows[2].cells[3], args.task_type)

    registration_paragraph = paragraph_after(document, info._element)
    replace_paragraph(
        registration_paragraph,
        f"{args.registration_number} — 최종 제출 시 접수번호를 공식 파일명에 반영",
        size=Pt(9),
    )


def resolved_versions(components: list[dict], group: str, name: str) -> str:
    versions = {
        str(component.get("version") or "").strip()
        for component in components
        if str(component.get("group") or "").strip() == group
        and str(component.get("name") or "").strip() == name
        and not component_is_test_only(component)
        and str(component.get("version") or "").strip()
    }

    def version_key(value: str) -> tuple:
        return tuple(
            (0, int(part)) if part.isdigit() else (1, part.lower())
            for part in re.split(r"([0-9]+)", value)
            if part
        )

    if not versions:
        return f"[확인 필요: {group}:{name} 해결 버전]"
    return "·".join(sorted(versions, key=version_key))


def component_is_test_only(component: dict) -> bool:
    return any(
        str(item.get("name") or "") == "cdx:maven:package:test"
        and str(item.get("value") or "").lower() == "true"
        for item in component.get("properties") or []
    )


def fill_report_body(document: Document, args: argparse.Namespace, components: list[dict]) -> None:
    table = document.tables[2]

    set_cell_text(table.rows[1].cells[1], "Reqover", bold=True)
    set_cell_text(table.rows[2].cells[1], "https://github.com/reqover-labs/reqover")
    set_cell_text(table.rows[3].cells[1], args.video_url)
    set_cell_text(
        table.rows[4].cells[1],
        "Spring MVC와 WebFlux에서 각 HTTP 요청이 실제로 실행한 메서드를 Java Agent로 자동 계측하여, "
        "endpoint-to-code 및 code-to-endpoint 관계를 JSON/HTML로 제공하는 오픈소스 개발자 도구이다.",
    )

    background = table.rows[6].cells[1]
    clear_cell(background)
    append_labeled_paragraph(
        background,
        "문제",
        "일반적인 집계 커버리지는 코드가 실행됐다는 사실은 보여주지만, 어떤 HTTP 요청이 해당 코드를 실행했는지는 기본 차원으로 제공하지 않는다. "
        "공유 service·validator가 많은 백엔드에서는 코드 변경 후 우선 재검증할 API를 경험과 넓은 회귀 테스트에 의존하기 쉽다.",
        first=True,
    )
    append_labeled_paragraph(
        background,
        "목표",
        "실제 요청 중 관측된 실행 관계를 endpoint별로 분리하고 역조회하여, 개발·QA·staging 환경의 디버깅과 회귀 테스트 우선순위 결정을 돕는다.",
    )
    append_labeled_paragraph(
        background,
        "해석 범위",
        "Reqover가 제시하는 관계는 관측된 실행의 하한이며, 보이지 않은 관계가 존재하지 않음을 보장하는 완전한 정적 영향 분석은 아니다.",
    )

    environment = table.rows[7].cells[1]
    clear_cell(environment)
    append_labeled_paragraph(environment, "언어·빌드", "Java 17 bytecode, JDK 17·21, Gradle 9.5.1", first=True)
    append_labeled_paragraph(
        environment,
        "프레임워크",
        f"Spring Boot {resolved_versions(components, 'org.springframework.boot', 'spring-boot')}, "
        f"Spring MVC {resolved_versions(components, 'org.springframework', 'spring-webmvc')}, "
        f"Spring WebFlux {resolved_versions(components, 'org.springframework', 'spring-webflux')}",
    )
    append_labeled_paragraph(
        environment,
        "핵심 기술",
        f"ASM {resolved_versions(components, 'org.ow2.asm', 'asm')}, "
        f"Reactor {resolved_versions(components, 'io.projectreactor', 'reactor-core')}, "
        f"Micrometer Context Propagation {resolved_versions(components, 'io.micrometer', 'context-propagation')}",
    )
    append_labeled_paragraph(
        environment,
        "검증",
        "JUnit 5.12.2, 별도 JVM Agent E2E, "
        "GitHub Actions Ubuntu/Temurin 17·21 matrix",
    )
    append_labeled_paragraph(environment, "개발 장비", args.development_environment)

    architecture = table.rows[8].cells[1]
    clear_cell(architecture)
    append_labeled_paragraph(
        architecture,
        "핵심 흐름",
        "HTTP 요청 → MVC Interceptor/WebFlux Filter → 요청 bucket 생성 → ASM Java Agent가 삽입한 method-entry probe → "
        "현재 요청 context로 hit 귀속 → snapshot 저장 → 양방향 JSON/HTML 리포트",
        first=True,
    )
    append_labeled_paragraph(architecture, "Core", "요청 수명주기, context, probe registry, in-memory 저장")
    append_labeled_paragraph(architecture, "Instrumentation/Agent", "선택한 애플리케이션 class를 로드할 때 method entry에 ReqoverProbe.hit을 삽입")
    append_labeled_paragraph(architecture, "Spring adapters", "MVC는 요청 ThreadLocal을, WebFlux는 Reactor Context와 context propagation을 사용")
    append_labeled_paragraph(architecture, "Report", "endpoint별 관측 메서드와 특정 코드가 관측된 endpoint를 양방향으로 집계")

    features = table.rows[9].cells[1]
    clear_cell(features)
    append_labeled_paragraph(
        features,
        "1. 요청 분리",
        "동시 실행된 GET /orders/{id}와 POST /payments를 별도 bucket으로 기록하고, 공통 SharedValidator만 양쪽 관계에 표시한다.",
        first=True,
    )
    add_picture(features, ASSET_MVC, "그림 1. MVC 요청별 실행 경로 분리 결과")
    append_labeled_paragraph(
        features,
        "2. 코드 역조회",
        "특정 class·method를 실행한 관측 endpoint를 찾아 코드 변경 후 우선 재검증할 API 후보를 좁힌다.",
    )
    add_picture(features, ASSET_REVERSE, "그림 2. code-to-endpoint 역조회 결과")
    append_labeled_paragraph(
        features,
        "3. 자동 계측",
        "사용자 코드에 수동 probe 호출을 넣지 않고 -javaagent 옵션으로 선택 package의 method entry를 계측한다.",
    )
    append_labeled_paragraph(
        features,
        "4. WebFlux thread-hop",
        "boundedElastic·parallel 등 thread가 바뀌는 표준 Reactor chain에서도 같은 요청 bucket으로 귀속한다.",
    )
    add_picture(features, ASSET_WEBFLUX, "그림 3. WebFlux thread 전환 후 동일 요청 귀속 결과")
    append_labeled_paragraph(
        features,
        "구동",
        "JDK 17 이상 환경에서 ./gradlew clean test 실행 후 ./scripts/run-agent-demo.sh mvc 8080 또는 webflux 8080을 실행하고, "
        "http://127.0.0.1:8080/reqover/report.html을 연다. 데모는 report endpoint를 loopback에만 노출한다.",
    )
    append_labeled_paragraph(features, "최종 검증", args.verification)

    effects = table.rows[10].cells[1]
    clear_cell(effects)
    append_labeled_paragraph(effects, "회귀 테스트", "관측 근거로 우선 재검증할 endpoint 후보를 좁혀 테스트 범위 결정 과정을 보조한다.", first=True)
    append_labeled_paragraph(effects, "디버깅·QA", "API 요청과 실제 실행 코드를 연결해 공유 로직과 reactive 실행 경로를 빠르게 이해한다.")
    append_labeled_paragraph(effects, "확장성", "향후 HTTP 외 메시지·배치·테스트 단위, CI 결과, persistent backend로 attribution 단위를 확장할 수 있다.")
    append_labeled_paragraph(effects, "오픈소스", "Apache-2.0, 재현 가능한 sample, 기여·보안 가이드와 SBOM을 제공해 외부 검증과 기여 기반을 마련한다.")

    other = table.rows[11].cells[1]
    clear_cell(other)
    append_labeled_paragraph(
        other,
        "혁신성·차별성",
        "집계 커버리지의 실행 여부를 요청 차원으로 확장하고, Java bytecode 계측과 WebFlux context propagation을 결합한다. "
        "endpoint→code와 code→endpoint를 함께 제공하며 JaCoCo를 대체하지 않고 요청 귀속 정보를 보완한다.",
        first=True,
    )
    append_labeled_paragraph(
        other,
        "현재 한계",
        "method-entry 수준, in-memory 저장, 명시적 include·불변 runtime exclude 정책, 인증 없는 sample report, 표준 Reactor chain 중심 검증 단계이다.",
    )
    append_labeled_paragraph(
        other,
        "로드맵",
        "재현 가능한 release artifact와 Gradle/Maven 연동, persistent storage와 인증, report filtering/export, JaCoCo XML 상호운용, "
        "반복 부하 측정을 순차적으로 추진한다.",
    )
    append_labeled_paragraph(other, "팀 역할·기여", args.team_contributions)
    append_labeled_paragraph(
        other,
        "개발 보조도구",
        "상용 생성형 AI(Codex)는 일부 코드·테스트·문서의 초안 작성 및 검토 보조에 활용했으며, 팀이 요구사항을 결정하고 결과를 검증·수정하여 "
        "최종 반영했다. 출품작에는 AI 모델이나 외부 추론 API가 포함되지 않는다.",
    )
    append_labeled_paragraph(
        other,
        "소감",
        "요청 수명주기와 reactive context를 bytecode 계측 결과에 안전하게 연결하면서 오귀속보다 미귀속을 우선하는 원칙의 중요성을 확인했다. "
        "기능 구현뿐 아니라 라이선스, SBOM, 재현 문서가 오픈소스 완성도의 일부임을 학습했다.",
    )
    append_labeled_paragraph(other, "중복수혜 확인", args.duplicate_benefit)


def component_license(component: dict) -> str:
    values: list[str] = []
    for entry in component.get("licenses") or []:
        if expression := entry.get("expression"):
            values.append(str(expression))
            continue
        license_info = entry.get("license") or {}
        value = license_info.get("id") or license_info.get("name")
        if value:
            values.append(str(value))
    unique = list(dict.fromkeys(values))
    return " OR ".join(unique) if unique else "[확인 필요: 라이선스 확인]"


def component_url(component: dict) -> str:
    references = component.get("externalReferences") or []
    priority = {"vcs": 0, "website": 1, "distribution": 2, "documentation": 3}
    for reference in sorted(references, key=lambda item: priority.get(item.get("type"), 99)):
        if reference.get("url"):
            url = str(reference["url"])
            replacements = {
                "scm:git:git://github.com/": "https://github.com/",
                "scm:git:https://github.com/": "https://github.com/",
                "git://github.com/": "https://github.com/",
                "http://github.com/": "https://github.com/",
                "http://developer.android.com/": "https://developer.android.com/",
            }
            for prefix, replacement in replacements.items():
                if url.startswith(prefix):
                    url = replacement + url.removeprefix(prefix)
                    break
            return url

    group = str(component.get("group") or "").strip()
    name = str(component.get("name") or "").strip()
    version = str(component.get("version") or "").strip()
    if group and name:
        base = f"https://central.sonatype.com/artifact/{quote(group, safe='.')}/{quote(name, safe='')}"
        return f"{base}/{quote(version, safe='.-_')}" if version else base
    if purl := component.get("purl"):
        return str(purl)
    return "[확인 필요: 공식 저장소 URL]"


def component_purpose(component: dict) -> str:
    group = str(component.get("group") or "").lower()
    name = str(component.get("name") or "").lower()
    key = f"{group}:{name}"

    if component_is_test_only(component):
        if "junit" in key or name in {"assertj-core", "awaitility", "hamcrest"}:
            return "단위·통합·E2E 테스트"
        if name.startswith("mockito-") or name in {"byte-buddy", "byte-buddy-agent", "objenesis"}:
            return "테스트 mocking 전이 의존성"
        return "테스트 전이 의존성"

    if name == "asm" or "ow2.asm" in group:
        return "Java bytecode method-entry 계측"
    if "spring-boot" in name:
        return "Spring Boot 실행·자동 설정 및 sample"
    if name in {"spring-webmvc", "spring-webflux"}:
        return "Spring MVC/WebFlux 요청 adapter"
    if name.startswith("spring-"):
        return "Spring Framework 전이 의존성"
    if "context-propagation" in name:
        return "Reactor Context와 ThreadLocal 전달"
    if name.startswith("reactor-core"):
        return "WebFlux reactive 실행과 thread-hop"
    if name.startswith("reactor-netty") or name.startswith("netty-"):
        return "WebFlux sample 네트워크 런타임"
    if name.startswith("tomcat-"):
        return "MVC sample 내장 HTTP 서버"
    if name.startswith("jackson-") or "json" in name:
        return "JSON 직렬화·테스트 지원"
    if "slf4j" in name or "logback" in name or "log4j" in name:
        return "로깅 API·구현체"
    if name.startswith("micrometer-"):
        return "Spring 관측·context 전이 의존성"
    if name.startswith("jakarta.") or group.startswith("jakarta."):
        return "Spring adapter 표준 Jakarta API"
    if name == "snakeyaml":
        return "Spring 설정 YAML 처리"
    if name == "reactive-streams":
        return "Reactive Streams 표준 API"
    return "빌드·실행·테스트 전이 의존성"


def load_sbom_components(path: Path) -> list[dict]:
    if not path.is_file():
        raise FileNotFoundError(
            f"CycloneDX SBOM not found: {path}\n"
            "Generate it with ./gradlew cyclonedxBom, preserve the final JSON under sbom/, or pass --sbom PATH."
        )
    with path.open("r", encoding="utf-8") as handle:
        data = json.load(handle)
    if data.get("bomFormat") != "CycloneDX":
        raise ValueError(f"Expected CycloneDX JSON SBOM: {path}")

    unique: dict[tuple[str, str, str], dict] = {}
    for component in data.get("components") or []:
        group = str(component.get("group") or "").strip()
        name = str(component.get("name") or "").strip()
        version = str(component.get("version") or "").strip()
        if not name or group == "io.reqover":
            continue
        unique[(group, name, version)] = component
    return [unique[key] for key in sorted(unique, key=lambda value: tuple(part.lower() for part in value))]


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def fill_sbom(document: Document, components: list[dict]) -> int:
    table = document.tables[4]
    repeat_table_header(table.rows[0])

    template_row = deepcopy(table.rows[1]._tr)
    for row in list(table.rows)[1:]:
        table._tbl.remove(row._tr)

    for index, component in enumerate(components, start=1):
        new_tr = deepcopy(template_row)
        table._tbl.append(new_tr)
        row = table.rows[-1]
        keep_row_together(row)
        values = [
            str(index),
            f"{component.get('group') or ''}:{component.get('name') or ''}".lstrip(":"),
            str(component.get("version") or "[확인 필요: 버전 확인]"),
            component_license(component),
            component_url(component),
            component_purpose(component),
        ]
        for column, value in enumerate(values):
            set_cell_text(
                row.cells[column],
                value,
                size=SBOM_FONT_SIZE,
                center=column in {0, 2, 3},
            )

    if not components:
        new_tr = deepcopy(template_row)
        table._tbl.append(new_tr)
        row = table.rows[-1]
        for column, value in enumerate(
            [
                "1",
                "[확인 필요: SBOM component 없음]",
                "",
                "",
                "",
                "CycloneDX 생성 설정과 대상 configuration 확인",
            ]
        ):
            set_cell_text(row.cells[column], value, size=SBOM_FONT_SIZE, center=column in {0, 2, 3})

    # The template note is not evidence and its required post-table paragraph can
    # become a blank trailing page after the table expands across many pages.
    note = next(paragraph for paragraph in document.paragraphs if "필요 시, 행을 추가" in paragraph.text)
    remove_element(note._p)
    return len(components)


def normalize_document_fonts(document: Document) -> None:
    def format_runs(paragraphs) -> None:
        for paragraph in paragraphs:
            for run in paragraph.runs:
                current_size = run.font.size or BODY_FONT_SIZE
                set_run_font(
                    run,
                    size=current_size,
                    bold=run.bold,
                    italic=run.italic,
                    color=(
                        str(run.font.color.rgb)
                        if run.font.color is not None and run.font.color.type is not None and run.font.color.rgb is not None
                        else "000000"
                    ),
                )

    format_runs(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                format_runs(cell.paragraphs)


def list_placeholders(document: Document) -> list[str]:
    values: list[str] = []

    def inspect_paragraphs(paragraphs) -> None:
        for paragraph in paragraphs:
            for match in re.findall(r"\[확인 필요:[^\]]+\]", paragraph.text):
                values.append(match)

    inspect_paragraphs(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                inspect_paragraphs(cell.paragraphs)
    return list(dict.fromkeys(values))


def build_document(args: argparse.Namespace) -> tuple[Path, int, list[str]]:
    template = args.template.resolve()
    sbom = args.sbom.resolve()
    output = args.output.resolve()
    verify_template(template)
    components = load_sbom_components(sbom)

    document = Document(str(template))
    remove_guide_and_ai_appendix(document)
    fill_basic_information(document, args)
    fill_report_body(document, args, components)
    component_count = fill_sbom(document, components)
    normalize_document_fonts(document)

    document.core_properties.title = "2026 오픈소스 개발자대회 결과보고서 - Reqover"
    document.core_properties.subject = "Reqover 요청 단위 런타임 커버리지 귀속 도구"
    document.core_properties.keywords = "Reqover, Java Agent, Spring MVC, WebFlux, coverage attribution"
    document.core_properties.author = "Reqover"
    document.core_properties.last_modified_by = "Reqover"

    placeholders = list_placeholders(document)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))
    if sha256(template) != TEMPLATE_SHA256:
        raise RuntimeError("Retained official template changed during generation")
    return output, component_count, placeholders


def main() -> None:
    args = parse_args()
    output, component_count, placeholders = build_document(args)
    print(f"Generated: {output}")
    print(f"SBOM components: {component_count}")
    if placeholders:
        print("Unresolved placeholders:")
        for placeholder in placeholders:
            print(f"- {placeholder}")


if __name__ == "__main__":
    main()
